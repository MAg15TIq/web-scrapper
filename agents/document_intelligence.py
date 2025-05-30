"""
Document Intelligence Agent for the self-aware web scraping system.

This agent handles all document types (PDF, DOC, DOCX, XLS, PPT, etc.).
"""
import asyncio
import logging
import time
import re
import json
import os
import io
from typing import Dict, Any, List, Optional, Set, Tuple, Union
import uuid
from urllib.parse import urlparse
import httpx

from agents.base import Agent
from models.message import Message, TaskMessage, ResultMessage, ErrorMessage, StatusMessage, Priority
from models.task import Task, TaskStatus, TaskType
from models.intelligence import (
    ContentType, InputType, WebsiteType, AgentCapability,
    AgentProfile, ContentAnalysisResult
)

# Try to import document processing libraries
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


class DocumentIntelligenceAgent(Agent):
    """
    Document Intelligence Agent that handles all document types.
    """
    def __init__(self, agent_id: Optional[str] = None, coordinator_id: str = "coordinator"):
        """
        Initialize a new Document Intelligence Agent.

        Args:
            agent_id: Unique identifier for the agent. If None, a UUID will be generated.
            coordinator_id: ID of the coordinator agent. Used for message routing.
        """
        super().__init__(agent_id=agent_id, agent_type="document_intelligence", coordinator_id=coordinator_id)

        # HTTP client for making requests
        self.client = httpx.AsyncClient(
            timeout=60.0,  # Longer timeout for document downloads
            follow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"}
        )

        # Check available libraries
        self.capabilities = {
            "pdf_basic": PYPDF2_AVAILABLE,
            "pdf_advanced": PDFMINER_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "excel": OPENPYXL_AVAILABLE,
            "data_processing": PANDAS_AVAILABLE
        }

        # Log available capabilities
        self.logger.info(f"Document processing capabilities: {self.capabilities}")

        # Cache for document processing results
        self.document_cache: Dict[str, Dict[str, Any]] = {}

        # Register message handlers
        self.register_handler("process_document", self._handle_process_document)
        self.register_handler("extract_text", self._handle_extract_text)
        self.register_handler("extract_tables", self._handle_extract_tables)
        self.register_handler("extract_metadata", self._handle_extract_metadata)

        # Start periodic tasks
        self._start_periodic_tasks()

    def _start_periodic_tasks(self) -> None:
        """Start periodic tasks for the Document Intelligence Agent."""
        asyncio.create_task(self._periodic_cache_cleanup())

    async def _periodic_cache_cleanup(self) -> None:
        """Periodically clean up the document cache."""
        while self.running:
            self.logger.debug("Running periodic cache cleanup")

            # Remove cache entries older than 1 hour
            current_time = time.time()
            expired_keys = []

            for key, cache_entry in self.document_cache.items():
                if current_time - cache_entry.get("timestamp", 0) > 3600:
                    expired_keys.append(key)

            # Remove expired entries
            for key in expired_keys:
                del self.document_cache[key]

            self.logger.debug(f"Removed {len(expired_keys)} expired cache entries")

            # Sleep for 15 minutes
            await asyncio.sleep(900)

    async def process_document(self, document_data: Union[str, bytes], document_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Process a document and extract its content.

        Args:
            document_data: The document data. Can be a URL, file path, or raw content.
            document_type: The document type if known. If None, it will be detected.

        Returns:
            A dictionary containing the processing results.
        """
        self.logger.info(f"Processing document: {type(document_data)}")

        # Generate a cache key
        if isinstance(document_data, str) and document_data.startswith(("http://", "https://")):
            cache_key = document_data  # Use URL as key
        elif isinstance(document_data, str) and (os.path.exists(document_data) or "/" in document_data or "\\" in document_data):
            cache_key = document_data  # Use file path as key
        elif isinstance(document_data, bytes):
            cache_key = str(hash(document_data))
        else:
            cache_key = str(uuid.uuid4())

        # Check cache first
        if cache_key in self.document_cache:
            self.logger.info(f"Using cached document processing result")
            return self.document_cache[cache_key]["result"]

        # Get document content if it's a URL or file path
        content = document_data
        if isinstance(document_data, str):
            if document_data.startswith(("http://", "https://")):
                # It's a URL, download the document
                content = await self._download_document(document_data)
            elif os.path.exists(document_data) or "/" in document_data or "\\" in document_data:
                # It's a file path, read the document
                try:
                    with open(document_data, "rb") as f:
                        content = f.read()
                except Exception as e:
                    self.logger.error(f"Error reading document file: {str(e)}")
                    content = b""

        # Determine document type if not provided
        if document_type is None:
            document_type = self._detect_document_type(content)

        # Process document based on type
        if document_type == "pdf":
            result = await self._process_pdf(content)
        elif document_type == "docx":
            result = await self._process_docx(content)
        elif document_type == "doc":
            result = await self._process_doc(content)
        elif document_type == "xlsx":
            result = await self._process_xlsx(content)
        elif document_type == "xls":
            result = await self._process_xls(content)
        elif document_type == "csv":
            result = await self._process_csv(content)
        else:
            result = {
                "error": f"Unsupported document type: {document_type}",
                "document_type": document_type
            }

        # Add document type to result
        result["document_type"] = document_type

        # Cache the result
        self.document_cache[cache_key] = {
            "result": result,
            "timestamp": time.time()
        }

        self.logger.info(f"Document processing complete: {document_type}")
        return result

    async def _download_document(self, url: str) -> bytes:
        """
        Download a document from a URL.

        Args:
            url: The URL to download from.

        Returns:
            The document content as bytes.
        """
        self.logger.info(f"Downloading document from URL: {url}")

        try:
            response = await self.client.get(url)
            return response.content
        except Exception as e:
            self.logger.error(f"Error downloading document: {str(e)}")
            return b""

    def _detect_document_type(self, content: Union[str, bytes]) -> str:
        """
        Detect the document type from its content.

        Args:
            content: The document content.

        Returns:
            The detected document type.
        """
        if isinstance(content, str):
            # If it's a URL, check the extension
            if content.startswith(("http://", "https://")):
                url_path = urlparse(content).path.lower()
                if url_path.endswith(".pdf"):
                    return "pdf"
                elif url_path.endswith(".docx"):
                    return "docx"
                elif url_path.endswith(".doc"):
                    return "doc"
                elif url_path.endswith(".xlsx"):
                    return "xlsx"
                elif url_path.endswith(".xls"):
                    return "xls"
                elif url_path.endswith(".csv"):
                    return "csv"
                elif url_path.endswith(".txt"):
                    return "txt"

            # If it's a file path, check the extension
            if "/" in content or "\\" in content:
                file_path = content.lower()
                if file_path.endswith(".pdf"):
                    return "pdf"
                elif file_path.endswith(".docx"):
                    return "docx"
                elif file_path.endswith(".doc"):
                    return "doc"
                elif file_path.endswith(".xlsx"):
                    return "xlsx"
                elif file_path.endswith(".xls"):
                    return "xls"
                elif file_path.endswith(".csv"):
                    return "csv"
                elif file_path.endswith(".txt"):
                    return "txt"

            # Otherwise, assume it's text
            return "txt"

        # If it's bytes, check for magic numbers
        if isinstance(content, bytes):
            # Check for PDF
            if content.startswith(b"%PDF-"):
                return "pdf"

            # Check for DOC
            if content.startswith(b"\xD0\xCF\x11\xE0"):
                return "doc"

            # Check for DOCX/XLSX (both are ZIP-based)
            if content.startswith(b"PK"):
                # Need more sophisticated checks to distinguish between DOCX and XLSX
                # For now, default to DOCX
                return "docx"

        # Default to unknown
        return "unknown"

    async def _process_pdf(self, content: bytes) -> Dict[str, Any]:
        """
        Process a PDF document.

        Args:
            content: The PDF content as bytes.

        Returns:
            A dictionary containing the processing results.
        """
        result = {
            "text": "",
            "metadata": {},
            "pages": [],
            "tables": []
        }

        # Check if we have PDF processing capabilities
        if not PYPDF2_AVAILABLE and not PDFMINER_AVAILABLE:
            result["error"] = "No PDF processing libraries available"
            return result

        # Create a file-like object from bytes
        pdf_file = io.BytesIO(content)

        # Extract text and metadata using PyPDF2
        if PYPDF2_AVAILABLE:
            try:
                pdf_reader = PyPDF2.PdfReader(pdf_file)

                # Extract metadata
                if pdf_reader.metadata:
                    result["metadata"] = {
                        "title": pdf_reader.metadata.get("/Title"),
                        "author": pdf_reader.metadata.get("/Author"),
                        "subject": pdf_reader.metadata.get("/Subject"),
                        "creator": pdf_reader.metadata.get("/Creator"),
                        "producer": pdf_reader.metadata.get("/Producer"),
                        "creation_date": pdf_reader.metadata.get("/CreationDate")
                    }

                # Extract text from each page
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    result["pages"].append({
                        "page_number": i + 1,
                        "text": page_text
                    })

                # Combine all pages
                result["text"] = "\n".join(page["text"] for page in result["pages"])

                # Add page count
                result["page_count"] = len(pdf_reader.pages)
            except Exception as e:
                self.logger.error(f"Error processing PDF with PyPDF2: {str(e)}")
                result["error"] = f"PyPDF2 error: {str(e)}"

        # If PyPDF2 failed or didn't extract text, try pdfminer
        if (not result["text"] or "error" in result) and PDFMINER_AVAILABLE:
            try:
                # Reset file pointer
                pdf_file.seek(0)

                # Extract text using pdfminer
                text = pdf_extract_text(pdf_file)
                result["text"] = text

                # If we didn't get pages from PyPDF2, split by page breaks
                if not result["pages"]:
                    # This is a simple heuristic, not perfect
                    pages = text.split("\f")
                    result["pages"] = [
                        {"page_number": i + 1, "text": page.strip()}
                        for i, page in enumerate(pages) if page.strip()
                    ]

                # Add page count if not already added
                if "page_count" not in result:
                    result["page_count"] = len(result["pages"])

                # Remove error if we succeeded with pdfminer
                if "error" in result:
                    del result["error"]
            except Exception as e:
                self.logger.error(f"Error processing PDF with pdfminer: {str(e)}")
                if "error" not in result:
                    result["error"] = f"pdfminer error: {str(e)}"

        # Extract tables (placeholder - would need a specialized library like tabula-py)
        result["tables"] = []

        return result

    async def _process_docx(self, content: bytes) -> Dict[str, Any]:
        """
        Process a DOCX document.

        Args:
            content: The DOCX content as bytes.

        Returns:
            A dictionary containing the processing results.
        """
        result = {
            "text": "",
            "metadata": {},
            "paragraphs": [],
            "tables": []
        }

        # Check if we have DOCX processing capabilities
        if not DOCX_AVAILABLE:
            result["error"] = "No DOCX processing libraries available"
            return result

        # Create a file-like object from bytes
        docx_file = io.BytesIO(content)

        try:
            # Open the document
            doc = docx.Document(docx_file)

            # Extract metadata
            core_properties = doc.core_properties
            result["metadata"] = {
                "title": core_properties.title,
                "author": core_properties.author,
                "subject": core_properties.subject,
                "keywords": core_properties.keywords,
                "created": str(core_properties.created) if core_properties.created else None,
                "modified": str(core_properties.modified) if core_properties.modified else None,
                "last_modified_by": core_properties.last_modified_by
            }

            # Extract paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    result["paragraphs"].append({
                        "index": i,
                        "text": para.text,
                        "style": para.style.name if para.style else None
                    })

            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)

                result["tables"].append({
                    "index": i,
                    "data": table_data
                })

            # Combine all paragraphs
            result["text"] = "\n".join(para["text"] for para in result["paragraphs"])

            # Add paragraph and table counts
            result["paragraph_count"] = len(result["paragraphs"])
            result["table_count"] = len(result["tables"])
        except Exception as e:
            self.logger.error(f"Error processing DOCX: {str(e)}")
            result["error"] = f"DOCX processing error: {str(e)}"

        return result

    async def _process_doc(self, content: bytes) -> Dict[str, Any]:
        """
        Process a DOC document.

        Args:
            content: The DOC content as bytes.

        Returns:
            A dictionary containing the processing results.
        """
        # DOC processing is more complex and requires external libraries like antiword or textract
        # For now, return a placeholder result
        return {
            "text": "",
            "metadata": {},
            "error": "DOC processing not implemented. Convert to DOCX for better support."
        }

    async def _process_xlsx(self, content: bytes) -> Dict[str, Any]:
        """
        Process an XLSX document.

        Args:
            content: The XLSX content as bytes.

        Returns:
            A dictionary containing the processing results.
        """
        result = {
            "sheets": [],
            "metadata": {}
        }

        # Check if we have Excel processing capabilities
        if not OPENPYXL_AVAILABLE:
            result["error"] = "No Excel processing libraries available"
            return result

        # Create a file-like object from bytes
        xlsx_file = io.BytesIO(content)

        try:
            # Open the workbook
            workbook = openpyxl.load_workbook(xlsx_file, read_only=True, data_only=True)

            # Extract metadata
            result["metadata"] = {
                "sheet_names": workbook.sheetnames,
                "sheet_count": len(workbook.sheetnames)
            }

            # Process each sheet
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]

                # Get sheet dimensions
                max_row = sheet.max_row
                max_col = sheet.max_column

                # Extract data
                sheet_data = []
                for row in range(1, min(max_row + 1, 1001)):  # Limit to 1000 rows for performance
                    row_data = []
                    for col in range(1, min(max_col + 1, 101)):  # Limit to 100 columns for performance
                        cell = sheet.cell(row=row, column=col)
                        row_data.append(str(cell.value) if cell.value is not None else "")
                    sheet_data.append(row_data)

                # Add sheet to result
                result["sheets"].append({
                    "name": sheet_name,
                    "data": sheet_data,
                    "row_count": max_row,
                    "column_count": max_col
                })

            # Convert to pandas DataFrames if available
            if PANDAS_AVAILABLE:
                for i, sheet in enumerate(result["sheets"]):
                    try:
                        # Create DataFrame
                        df = pd.DataFrame(sheet["data"][1:], columns=sheet["data"][0] if sheet["data"] else [])

                        # Add basic statistics
                        numeric_columns = df.select_dtypes(include=["number"]).columns
                        if not numeric_columns.empty:
                            stats = df[numeric_columns].describe().to_dict()
                            result["sheets"][i]["statistics"] = stats
                    except Exception as e:
                        self.logger.error(f"Error creating DataFrame for sheet {sheet['name']}: {str(e)}")
        except Exception as e:
            self.logger.error(f"Error processing XLSX: {str(e)}")
            result["error"] = f"XLSX processing error: {str(e)}"

        return result

    async def _process_xls(self, content: bytes) -> Dict[str, Any]:
        """
        Process an XLS document.

        Args:
            content: The XLS content as bytes.

        Returns:
            A dictionary containing the processing results.
        """
        # XLS processing is more complex and requires additional libraries
        # For now, return a placeholder result
        return {
            "sheets": [],
            "metadata": {},
            "error": "XLS processing not implemented. Convert to XLSX for better support."
        }

    async def _process_csv(self, content: bytes) -> Dict[str, Any]:
        """
        Process a CSV document.

        Args:
            content: The CSV content as bytes.

        Returns:
            A dictionary containing the processing results.
        """
        result = {
            "data": [],
            "metadata": {}
        }

        # Try to decode the content
        try:
            text_content = content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text_content = content.decode("latin-1")
            except Exception:
                result["error"] = "Failed to decode CSV content"
                return result

        # Split into lines
        lines = text_content.split("\n")
        non_empty_lines = [line for line in lines if line.strip()]

        if not non_empty_lines:
            result["error"] = "CSV file is empty"
            return result

        # Detect delimiter
        first_line = non_empty_lines[0]
        comma_count = first_line.count(",")
        semicolon_count = first_line.count(";")
        tab_count = first_line.count("\t")

        if semicolon_count > comma_count and semicolon_count > tab_count:
            delimiter = ";"
        elif tab_count > comma_count and tab_count > semicolon_count:
            delimiter = "\t"
        else:
            delimiter = ","

        # Parse CSV
        for line in non_empty_lines:
            result["data"].append(line.split(delimiter))

        # Add metadata
        result["metadata"] = {
            "delimiter": delimiter,
            "row_count": len(result["data"]),
            "column_count": len(result["data"][0]) if result["data"] else 0,
            "has_header": True  # Assume first row is header
        }

        # Convert to pandas DataFrame if available
        if PANDAS_AVAILABLE:
            try:
                # Create DataFrame
                df = pd.DataFrame(result["data"][1:], columns=result["data"][0] if result["data"] else [])

                # Add basic statistics
                numeric_columns = df.select_dtypes(include=["number"]).columns
                if not numeric_columns.empty:
                    stats = df[numeric_columns].describe().to_dict()
                    result["statistics"] = stats
            except Exception as e:
                self.logger.error(f"Error creating DataFrame for CSV: {str(e)}")

        return result

    async def extract_text(self, document_data: Union[str, bytes], document_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract text from a document.

        Args:
            document_data: The document data. Can be a URL, file path, or raw content.
            document_type: The document type if known. If None, it will be detected.

        Returns:
            A dictionary containing the extracted text.
        """
        # Process the document
        result = await self.process_document(document_data, document_type)

        # Extract just the text-related information
        text_result = {
            "document_type": result.get("document_type", "unknown"),
            "text": result.get("text", ""),
            "page_count": result.get("page_count", 0),
            "paragraph_count": result.get("paragraph_count", 0)
        }

        # Add pages if available
        if "pages" in result:
            text_result["pages"] = result["pages"]

        # Add paragraphs if available
        if "paragraphs" in result:
            text_result["paragraphs"] = result["paragraphs"]

        # Add error if present
        if "error" in result:
            text_result["error"] = result["error"]

        return text_result

    async def extract_tables(self, document_data: Union[str, bytes], document_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract tables from a document.

        Args:
            document_data: The document data. Can be a URL, file path, or raw content.
            document_type: The document type if known. If None, it will be detected.

        Returns:
            A dictionary containing the extracted tables.
        """
        # Process the document
        result = await self.process_document(document_data, document_type)

        # Extract just the table-related information
        table_result = {
            "document_type": result.get("document_type", "unknown"),
            "tables": result.get("tables", []),
            "table_count": len(result.get("tables", []))
        }

        # Add sheets if available (for Excel files)
        if "sheets" in result:
            table_result["sheets"] = result["sheets"]
            table_result["sheet_count"] = len(result["sheets"])

        # Add error if present
        if "error" in result:
            table_result["error"] = result["error"]

        return table_result

    async def extract_metadata(self, document_data: Union[str, bytes], document_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract metadata from a document.

        Args:
            document_data: The document data. Can be a URL, file path, or raw content.
            document_type: The document type if known. If None, it will be detected.

        Returns:
            A dictionary containing the extracted metadata.
        """
        # Process the document
        result = await self.process_document(document_data, document_type)

        # Extract just the metadata-related information
        metadata_result = {
            "document_type": result.get("document_type", "unknown"),
            "metadata": result.get("metadata", {})
        }

        # Add page count if available
        if "page_count" in result:
            metadata_result["page_count"] = result["page_count"]

        # Add paragraph count if available
        if "paragraph_count" in result:
            metadata_result["paragraph_count"] = result["paragraph_count"]

        # Add table count if available
        if "table_count" in result:
            metadata_result["table_count"] = result["table_count"]

        # Add error if present
        if "error" in result:
            metadata_result["error"] = result["error"]

        return metadata_result

    async def _handle_process_document(self, message: Message) -> None:
        """
        Handle a process document message.

        Args:
            message: The message to handle.
        """
        if not hasattr(message, "document_data"):
            self.logger.warning("Received process_document message without document_data")
            return

        try:
            # Get document type if provided
            document_type = message.document_type if hasattr(message, "document_type") else None

            # Process the document
            result = await self.process_document(message.document_data, document_type)

            # Send result
            response = ResultMessage(
                sender_id=self.agent_id,
                recipient_id=message.sender_id,
                task_id=message.task_id if hasattr(message, "task_id") else None,
                result=result
            )
            self.outbox.put(response)
        except Exception as e:
            self.logger.error(f"Error processing document: {str(e)}")

            # Send error
            error_message = ErrorMessage(
                sender_id=self.agent_id,
                recipient_id=message.sender_id,
                task_id=message.task_id if hasattr(message, "task_id") else None,
                error=str(e)
            )
            self.outbox.put(error_message)

    async def _handle_extract_text(self, message: Message) -> None:
        """
        Handle an extract text message.

        Args:
            message: The message to handle.
        """
        if not hasattr(message, "document_data"):
            self.logger.warning("Received extract_text message without document_data")
            return

        try:
            # Get document type if provided
            document_type = message.document_type if hasattr(message, "document_type") else None

            # Extract text
            result = await self.extract_text(message.document_data, document_type)

            # Send result
            response = ResultMessage(
                sender_id=self.agent_id,
                recipient_id=message.sender_id,
                task_id=message.task_id if hasattr(message, "task_id") else None,
                result=result
            )
            self.outbox.put(response)
        except Exception as e:
            self.logger.error(f"Error extracting text: {str(e)}")

            # Send error
            error_message = ErrorMessage(
                sender_id=self.agent_id,
                recipient_id=message.sender_id,
                task_id=message.task_id if hasattr(message, "task_id") else None,
                error=str(e)
            )
            self.outbox.put(error_message)

    async def _handle_extract_tables(self, message: Message) -> None:
        """
        Handle an extract tables message.

        Args:
            message: The message to handle.
        """
        if not hasattr(message, "document_data"):
            self.logger.warning("Received extract_tables message without document_data")
            return

        try:
            # Get document type if provided
            document_type = message.document_type if hasattr(message, "document_type") else None

            # Extract tables
            result = await self.extract_tables(message.document_data, document_type)

            # Send result
            response = ResultMessage(
                sender_id=self.agent_id,
                recipient_id=message.sender_id,
                task_id=message.task_id if hasattr(message, "task_id") else None,
                result=result
            )
            self.outbox.put(response)
        except Exception as e:
            self.logger.error(f"Error extracting tables: {str(e)}")

            # Send error
            error_message = ErrorMessage(
                sender_id=self.agent_id,
                recipient_id=message.sender_id,
                task_id=message.task_id if hasattr(message, "task_id") else None,
                error=str(e)
            )
            self.outbox.put(error_message)

    async def _handle_extract_metadata(self, message: Message) -> None:
        """
        Handle an extract metadata message.

        Args:
            message: The message to handle.
        """
        if not hasattr(message, "document_data"):
            self.logger.warning("Received extract_metadata message without document_data")
            return

        try:
            # Get document type if provided
            document_type = message.document_type if hasattr(message, "document_type") else None

            # Extract metadata
            result = await self.extract_metadata(message.document_data, document_type)

            # Send result
            response = ResultMessage(
                sender_id=self.agent_id,
                recipient_id=message.sender_id,
                task_id=message.task_id if hasattr(message, "task_id") else None,
                result=result
            )
            self.outbox.put(response)
        except Exception as e:
            self.logger.error(f"Error extracting metadata: {str(e)}")

            # Send error
            error_message = ErrorMessage(
                sender_id=self.agent_id,
                recipient_id=message.sender_id,
                task_id=message.task_id if hasattr(message, "task_id") else None,
                error=str(e)
            )
            self.outbox.put(error_message)

    async def execute_task(self, task: Task) -> Dict[str, Any]:
        """
        Execute a task assigned to this agent.

        Args:
            task: The task to execute.

        Returns:
            A dictionary containing the result of the task.
        """
        self.logger.info(f"Executing task: {task.type}")

        if task.type == TaskType.PROCESS_DOCUMENT:
            # Process document
            document_data = task.parameters.get("document_data")
            if not document_data:
                raise ValueError("Missing document_data parameter")

            document_type = task.parameters.get("document_type")
            return await self.process_document(document_data, document_type)

        elif task.type == TaskType.EXTRACT_TEXT:
            # Extract text
            document_data = task.parameters.get("document_data")
            if not document_data:
                raise ValueError("Missing document_data parameter")

            document_type = task.parameters.get("document_type")
            return await self.extract_text(document_data, document_type)

        elif task.type == TaskType.EXTRACT_TABLES:
            # Extract tables
            document_data = task.parameters.get("document_data")
            if not document_data:
                raise ValueError("Missing document_data parameter")

            document_type = task.parameters.get("document_type")
            return await self.extract_tables(document_data, document_type)

        elif task.type == TaskType.EXTRACT_METADATA:
            # Extract metadata
            document_data = task.parameters.get("document_data")
            if not document_data:
                raise ValueError("Missing document_data parameter")

            document_type = task.parameters.get("document_type")
            return await self.extract_metadata(document_data, document_type)

        else:
            raise ValueError(f"Unsupported task type: {task.type}")
